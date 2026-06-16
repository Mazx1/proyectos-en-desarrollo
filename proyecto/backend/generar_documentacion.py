"""Genera documento Word con la explicacion del backend."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUTPUT = r"c:\Users\cazava\Desktop\Github Manuel Zarate\proyecto\backend\Documentacion_Backend.docx"

doc = Document()

# Titulo principal
title = doc.add_heading("Documentacion del Backend", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("API REST con autenticacion JWT y control de roles")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

doc.add_paragraph(f"Fecha: {date.today().strftime('%d de %B de %Y')}")
doc.add_paragraph("Proyecto: backend-auth-roles")
doc.add_paragraph()

# 1. Introduccion
doc.add_heading("1. Introduccion", level=1)
doc.add_paragraph(
    "Este documento describe la creacion y el funcionamiento del backend desarrollado "
    "para el proyecto. Se trata de una API REST construida con Node.js y Express, que "
    "gestiona el registro e inicio de sesion de usuarios, emite tokens JWT para "
    "autenticacion y controla el acceso a recursos segun el rol del usuario "
    "(usuario, admin o superadmin)."
)
doc.add_paragraph(
    "La base de datos utilizada es MongoDB, accedida mediante Mongoose como ODM "
    "(Object Document Mapper). El backend esta pensado para ser consumido por un "
    "frontend (por ejemplo, una aplicacion React) que envie peticiones HTTP con "
    "JSON."
)

# 2. Tecnologias
doc.add_heading("2. Tecnologias utilizadas", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Tecnologia"
hdr[1].text = "Proposito"
techs = [
    ("Node.js", "Entorno de ejecucion de JavaScript en el servidor"),
    ("Express 4", "Framework web para crear rutas y middleware"),
    ("MongoDB + Mongoose", "Base de datos NoSQL y modelado de esquemas"),
    ("bcryptjs", "Cifrado seguro de contrasenas (hash)"),
    ("jsonwebtoken (JWT)", "Generacion y validacion de tokens de sesion"),
    ("dotenv", "Carga de variables de entorno desde archivo .env"),
    ("cors", "Permite peticiones desde otros origenes (frontend)"),
    ("nodemon", "Reinicio automatico del servidor en desarrollo"),
]
for tech, purpose in techs:
    row = table.add_row().cells
    row[0].text = tech
    row[1].text = purpose

doc.add_paragraph()

# 3. Estructura
doc.add_heading("3. Estructura del proyecto", level=1)
doc.add_paragraph("El backend sigue una organizacion modular por responsabilidades:")
doc.add_paragraph("backend/", style="List Bullet")
doc.add_paragraph("  package.json          - Dependencias y scripts npm", style="List Bullet")
doc.add_paragraph("  .env.example          - Plantilla de variables de entorno", style="List Bullet")
doc.add_paragraph("  src/", style="List Bullet")
doc.add_paragraph("    server.js           - Punto de entrada: Express + MongoDB", style="List Bullet")
doc.add_paragraph("    models/User.js      - Esquema del usuario en MongoDB", style="List Bullet")
doc.add_paragraph("    routes/auth.js      - Rutas de autenticacion y roles", style="List Bullet")
doc.add_paragraph("    middleware/auth.js  - Verificacion JWT y permisos por rol", style="List Bullet")

# 4. Configuracion
doc.add_heading("4. Configuracion e instalacion", level=1)
doc.add_heading("4.1 Variables de entorno", level=2)
doc.add_paragraph(
    "Se crea un archivo .env en la carpeta backend (copiando .env.example) con las "
    "siguientes variables:"
)
env_table = doc.add_table(rows=1, cols=2)
env_table.style = "Table Grid"
env_table.rows[0].cells[0].text = "Variable"
env_table.rows[0].cells[1].text = "Descripcion"
env_vars = [
    ("PORT", "Puerto del servidor (por defecto 5000)"),
    ("MONGODB_URI", "Cadena de conexion a la base de datos MongoDB Atlas"),
    ("JWT_SECRET", "Clave secreta para firmar y verificar los tokens JWT"),
]
for var, desc in env_vars:
    r = env_table.add_row().cells
    r[0].text = var
    r[1].text = desc

doc.add_heading("4.2 Pasos de instalacion", level=2)
steps = [
    "Abrir una terminal en la carpeta backend del proyecto.",
    "Ejecutar: npm install (instala todas las dependencias del package.json).",
    "Crear el archivo .env con PORT, MONGODB_URI y JWT_SECRET.",
    "Para desarrollo: npm run dev (usa nodemon para recargar al guardar cambios).",
    "Para produccion: npm start (ejecuta node src/server.js).",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}", style="List Number")

# 5. Modelo de datos
doc.add_heading("5. Modelo de datos: Usuario", level=1)
doc.add_paragraph(
    "El archivo src/models/User.js define el esquema Mongoose para la coleccion "
    "users en MongoDB. Los campos son:"
)
user_table = doc.add_table(rows=1, cols=3)
user_table.style = "Table Grid"
user_table.rows[0].cells[0].text = "Campo"
user_table.rows[0].cells[1].text = "Tipo"
user_table.rows[0].cells[2].text = "Detalle"
user_fields = [
    ("name", "String", "Nombre del usuario (obligatorio)"),
    ("email", "String", "Correo unico, en minusculas (obligatorio)"),
    ("password", "String", "Contrasena hasheada, minimo 6 caracteres"),
    ("role", "String", "usuario | admin | superadmin (default: usuario)"),
    ("createdAt / updatedAt", "Date", "Marcas de tiempo automaticas"),
]
for f, t, d in user_fields:
    r = user_table.add_row().cells
    r[0].text = f
    r[1].text = t
    r[2].text = d

doc.add_paragraph(
    "Las contrasenas nunca se almacenan en texto plano: al registrarse, se aplica "
    "bcrypt.hash con 10 rondas de sal antes de guardar en la base de datos."
)

# 6. Servidor
doc.add_heading("6. Punto de entrada: server.js", level=1)
doc.add_paragraph("El archivo src/server.js realiza las siguientes tareas:")
server_tasks = [
    "Carga las variables de entorno con dotenv.config().",
    "Crea la aplicacion Express y habilita CORS y el parser JSON.",
    "Define la ruta raiz GET / que responde con estado de la API.",
    "Monta las rutas de autenticacion en /api/auth.",
    "Conecta a MongoDB con mongoose.connect(MONGODB_URI).",
    "Si la conexion falla, muestra error y termina el proceso.",
    "Si la conexion es exitosa, inicia el servidor en el puerto configurado.",
]
for t in server_tasks:
    doc.add_paragraph(t, style="List Bullet")

# 7. Rutas API
doc.add_heading("7. Endpoints de la API", level=1)
doc.add_paragraph("Todas las rutas de autenticacion estan bajo el prefijo /api/auth:")

api_table = doc.add_table(rows=1, cols=4)
api_table.style = "Table Grid"
api_table.rows[0].cells[0].text = "Metodo"
api_table.rows[0].cells[1].text = "Ruta"
api_table.rows[0].cells[2].text = "Auth"
api_table.rows[0].cells[3].text = "Descripcion"
endpoints = [
    ("GET", "/", "No", "Comprueba que la API esta activa"),
    ("POST", "/api/auth/register", "No", "Registra un nuevo usuario y devuelve JWT"),
    ("POST", "/api/auth/login", "No", "Inicia sesion con email y password"),
    ("GET", "/api/auth/me", "Si (JWT)", "Devuelve datos del usuario autenticado"),
    ("GET", "/api/auth/admin", "Si + rol admin/superadmin", "Ruta protegida para administradores"),
    ("GET", "/api/auth/superadmin", "Si + rol superadmin", "Ruta exclusiva para superadmin"),
]
for method, route, auth, desc in endpoints:
    r = api_table.add_row().cells
    r[0].text = method
    r[1].text = route
    r[2].text = auth
    r[3].text = desc

doc.add_heading("7.1 Registro (POST /api/auth/register)", level=2)
doc.add_paragraph("Cuerpo JSON esperado:")
p = doc.add_paragraph()
p.add_run('{\n  "name": "Juan Perez",\n  "email": "juan@ejemplo.com",\n  "password": "miClave123",\n  "role": "usuario"\n}').font.name = "Consolas"
doc.add_paragraph(
    "El campo role es opcional; si no se envia, se asigna usuario. Respuesta exitosa (201): "
    "mensaje, token JWT y objeto user sin la contrasena."
)

doc.add_heading("7.2 Login (POST /api/auth/login)", level=2)
doc.add_paragraph("Cuerpo JSON: email y password. Si las credenciales son correctas, "
    "devuelve token y datos del usuario. Si no, responde 401.")

doc.add_heading("7.3 Rutas protegidas", level=2)
doc.add_paragraph(
    "Las rutas /me, /admin y /superadmin requieren el header HTTP: "
    "Authorization: Bearer <token>. El middleware authRequired extrae y verifica "
    "el JWT. El middleware requireRoles comprueba que el rol del usuario este "
    "en la lista permitida; si no, responde 403 Forbidden."
)

# 8. Seguridad
doc.add_heading("8. Seguridad implementada", level=1)
security = [
    "Hash de contrasenas con bcrypt (10 salt rounds).",
    "Tokens JWT firmados con JWT_SECRET y expiracion de 1 dia.",
    "Validacion de token en cada peticion protegida.",
    "Control de acceso basado en roles (RBAC): usuario, admin, superadmin.",
    "Email unico en base de datos para evitar duplicados.",
    "Mensajes genericos en login (Credenciales invalidas) para no revelar si el email existe.",
]
for s in security:
    doc.add_paragraph(s, style="List Bullet")

# 9. Flujo
doc.add_heading("9. Flujo de autenticacion", level=1)
doc.add_paragraph(
    "1. El cliente envia POST /api/auth/register o /login con los datos.\n"
    "2. El servidor valida, busca o crea el usuario en MongoDB.\n"
    "3. En login, compara la contrasena con bcrypt.compare.\n"
    "4. Se genera un JWT con id, email, role y name del usuario.\n"
    "5. El cliente guarda el token (localStorage, cookie, etc.).\n"
    "6. En peticiones siguientes, envia Authorization: Bearer <token>.\n"
    "7. authRequired decodifica el token y adjunta req.user.\n"
    "8. requireRoles verifica el rol antes de permitir el acceso a rutas restringidas."
)

# 10. Codigos HTTP
doc.add_heading("10. Codigos de respuesta HTTP", level=1)
codes = [
    ("200", "Operacion exitosa (login, rutas protegidas)"),
    ("201", "Usuario registrado correctamente"),
    ("400", "Faltan campos requeridos en el body"),
    ("401", "Credenciales invalidas o token ausente/invalido"),
    ("403", "Usuario autenticado pero sin permisos para el recurso"),
    ("409", "Email ya registrado"),
    ("500", "Error interno del servidor"),
]
for code, meaning in codes:
    doc.add_paragraph(f"{code}: {meaning}", style="List Bullet")

# 11. Integracion frontend
doc.add_heading("11. Integracion con el frontend", level=1)
doc.add_paragraph(
    "El frontend debe configurar la URL base de la API (por ejemplo "
    "http://localhost:5000) y enviar peticiones fetch o axios con "
    "Content-Type: application/json. Tras login o registro, debe almacenar "
    "el token y enviarlo en el header Authorization en cada peticion a rutas "
    "protegidas. Segun el rol devuelto en user.role, la interfaz puede mostrar "
    "u ocultar secciones de administracion."
)

# 12. Conclusion
doc.add_heading("12. Conclusion", level=1)
doc.add_paragraph(
    "El backend creado proporciona una base solida y escalable para aplicaciones "
    "que requieren autenticacion de usuarios y distintos niveles de permisos. "
    "Su arquitectura modular (modelos, rutas, middleware) facilita anadir nuevas "
    "entidades y endpoints en el futuro, como gestion de productos, pedidos u "
    "otros modulos del negocio, reutilizando el mismo sistema de JWT y roles."
)

doc.save(OUTPUT)
print(f"Documento creado: {OUTPUT}")
